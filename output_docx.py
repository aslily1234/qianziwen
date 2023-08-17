from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from pypinyin import pinyin, lazy_pinyin, Style

#参考https://github.com/pwxcoo/chinese-xinhua ，导入成语信息
import json
file = 'data/idiom.json'
idiom_info_list = None
with open(file,'r', encoding='UTF-8') as f:
     idiom_info_list = json.load(f)

#成语-拼音索引
pinyin_dict = {} #{"阿鼻地狱" : "ā bí dì yù", }
'''
{
    "derivation": "语出《法华经·法师功德品》下至阿鼻地狱。”",
    "example": "但也有少数意志薄弱的……逐步上当，终至堕入～。★《上饶集中营·炼狱杂记》",
    "explanation": "阿鼻梵语的译音，意译为无间”，即痛苦无有间断之意。常用来比喻黑暗的社会和严酷的牢狱。又比喻无法摆脱的极其痛苦的境地。",
    "pinyin": "ā bí dì yù",
    "word": "阿鼻地狱",
    "abbreviation": "abdy"
}
'''
#构建拼音索引
for data in idiom_info_list:
    idiom = data['word']
    pinyin_dict[idiom] = data['pinyin']

#获取字体长度
from PIL import Image, ImageFont, ImageDraw
def get_text_length(text, font_size, font_path = "font/msyhbd.ttf"):
    font = ImageFont.truetype(font_path, font_size)
    image = Image.new('RGB', (1, 1))
    draw = ImageDraw.Draw(image)
    return draw.textlength(text, font=font)

'''
text = "yí zūn jiù jiào ， jiāo yī shí bǎi"
font_size = 12
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "移樽就教，教一识百"
font_size = 44
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "huí tiān yùn dǒu ， dǒu zhuǎn xīng yí"
font_size = 22
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "回天运斗，斗转星移"
font_size = 44
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")
'''

def add_title(document, title):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(72)
    run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

#将一行成语加入word中，并对其进行拼音标注
def add_text_paragraph(document, text, text_font_size=44):
    #获得文字对应的拼音(方式一，通过lazy_pinyin获取)
    pinyin_text = ' '.join(lazy_pinyin(text, style=Style.TONE))

    #获得文字对应的拼音(方式二，通过pinyin_dict查找)
    idiom_list = text.split('，')
    pinyin_text = pinyin_dict[idiom_list[0]] + '，' + pinyin_dict[idiom_list[1]]

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    #设置段后间距（单倍行距，段后间距为0）
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_after = Pt(0)

    #固定文字字号为text_font_size(默认为44），推算拼音的字号大小，使得拼音与文字的长度保持一致
    text_length = get_text_length(text, text_font_size)

    #寻找最佳字号，使得拼音的长度与文字的长度保持一致
    pinyin_font_size = int(text_font_size / 2)
    min_gap = text_length
    min_pinyin_font_size = pinyin_font_size
    for size in range(pinyin_font_size - 10, pinyin_font_size + 10, 1):
        pinyin_length = get_text_length(pinyin_text, size)

        if abs(pinyin_length - text_length) < min_gap:
            min_gap = abs(pinyin_length - text_length)
            min_pinyin_font_size = size

        # print(size, abs(pinyin_length - text_length))

    #使用间距最小的字号作为拼音的字号
    pinyin_font_size = min_pinyin_font_size

    #写入拼音行
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(pinyin_text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(pinyin_font_size)

    #写入文本行
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(text_font_size)

def jielong_list_to_docx(jielong, outputfile, title="成语接龙"):
    document = Document()

    title = '成语接龙'
    add_title(document, title)

    for i in range(1, len(jielong), 2):
        text = jielong[i-1] + '，' + jielong[i]
        add_text_paragraph(document, text)

        #处理奇数情况的最后一个成语
        if i == len(jielong) - 2:
            text = jielong[i+1]
            add_text_paragraph(document, text)

    document.save(outputfile)

if __name__ == '__main__':
    document = Document()

    title = '成语接龙'
    add_title(document, title)

    text = '字顺文从，从风而服，服低做小，小试锋芒'
    text = '字顺文从，从风而服'
    add_text_paragraph(document, text)

    outputfile = 'output3.docx'
    document.save(outputfile)

    jielong = ['回天运斗', '斗转星移', '移樽就教', '教一识百', '百足之虫', '虫鱼之学', '学以致用', '用智铺谋', '谋无遗策', '策无遗算']
    jielong_list_to_docx(jielong, "jielong.docx", title="成语接龙")
