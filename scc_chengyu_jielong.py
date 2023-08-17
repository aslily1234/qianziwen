
from collections import defaultdict

#初始化，构造字频索引，成语首字索引，成语首字同音索引，成语尾字索引，成语尾字同音索引，首尾同字相连版本成语接龙网状图，首尾同音字相连版本成语接龙网状图
import json 
from xpinyin import Pinyin
import random


p = Pinyin()

#参考https://github.com/pwxcoo/chinese-xinhua ，导入成语信息
file = 'data/idiom.json'
idiom_info_list = None
with open(file,'r', encoding='UTF-8') as f:
     idiom_info_list = json.load(f)

#设置低频词的阈值
min_word_freq = 5 

#字频索引
word_dict = {}
#成语首字索引
idiom_head_list = {}
#成语尾字索引
idiom_tail_list = {}
#成语-拼音索引
pinyin_dict = {} #{"阿鼻地狱" : "ā bí dì yù", }

#首尾同字相连版本成语接龙网状图
idiom_jielong_graph = {}


'''
{
    "derivation": "语出《法华经·法师功德品》下至阿鼻地狱。”",
    "example": "但也有少数意志薄弱的……逐步上当，终至堕入～。★《上饶集中营·炼狱杂记》",
    "explanation": "阿鼻梵语的译音，意译为无间”，即痛苦无有间断之意。常用来比喻黑暗的社会和严酷的牢狱。又比喻无法摆脱的极其痛苦的境地。",
    "pinyin": "ā bí dì yù",
    "word": "阿鼻地狱",
    "abbreviation": "abdy"
}
'''
#统计字频，构建拼音索引
for data in idiom_info_list:
    idiom = data['word']
    pinyin_dict[idiom] = data['pinyin']

    idiom_word_list = list(idiom)

    #统计字频
    for word in idiom_word_list:
        if word in word_dict:
            word_dict[word] += 1
        else:
            word_dict[word] = 1

#对字频库进行排序 key使用lambda匿名函数按键进行排序
sort_word_dict = sorted(word_dict.items(),key = lambda x:x[1], reverse=False)

#输出字频索引
from openpyxl import Workbook

def write_dict_to_excel(data_dict, filename):
    # 创建一个新的Excel工作簿
    workbook = Workbook()

    # 获取活动的工作表
    worksheet = workbook.active

    # 写入字典的键作为表头
    headers = list(data_dict.keys())
    worksheet.append(headers)

    # 写入字典的值
    values = list(data_dict.values())
    worksheet.append(values)

    # 保存Excel文件
    workbook.save(filename)

write_dict_to_excel(word_dict, 'word_dict.xlsx')


#过滤掉包含低频字的成语，构建成语首字索引，成语尾字索引
for data in idiom_info_list:
    idiom = data['word']

    idiom_word_list = list(idiom)

    #过滤掉非四字成语
    if len(idiom_word_list) != 4:
        print(idiom)
        continue

    #过滤掉有重复字的成语
    if len(set(idiom_word_list)) != 4:
        print(idiom)
        continue

    #过滤掉包含低频字的成语
    not_freq = 0
    for word in idiom_word_list:
        if word_dict[word] < min_word_freq:
            not_freq = 1
            break
    if not_freq == 1:
        continue

    #构建成语首字索引库
    first_word = idiom_word_list[0]

    if first_word in idiom_head_list:
        idiom_head_list[first_word].append(idiom)
    else:
        idiom_head_list[first_word] = [idiom]

    #构建成语尾字索引库
    last_word = idiom_word_list[-1]

    if last_word in idiom_tail_list:
        idiom_tail_list[last_word].append(idiom)
    else:
        idiom_tail_list[last_word] = [idiom]

#过滤掉包含低频字的成语，构建首尾同字相连版本成语接龙网状图
for data in idiom_info_list:
    idiom = data['word']

    idiom_word_list = list(idiom)

    #过滤掉非四字成语
    if len(idiom_word_list) != 4:
        print(idiom)
        continue

    #过滤掉有重复字的成语
    if len(set(idiom_word_list)) != 4:
        print(idiom)
        continue

    #过滤掉包含低频字的成语
    not_freq = 0
    for word in idiom_word_list:
        if word_dict[word] < min_word_freq:
            not_freq = 1
            break
    if not_freq == 1:
        continue

    last_word = idiom_word_list[-1]
    
    if idiom not in idiom_jielong_graph:
        #初始化该成语的邻接表内容
        idiom_jielong_graph[idiom] = list()

        if last_word in idiom_head_list:
            for key in idiom_head_list[last_word]:
                idiom_jielong_graph[idiom].append(key)

#输出统计信息
print('word_num:', len(word_dict.keys()))
print('idiom_head_num:', len(idiom_head_list.keys()))
print('idiom_tail_num:', len(idiom_tail_list.keys()))
print('idiom_jielong_graph:', len(idiom_jielong_graph.keys()))


'''
edges={
    "a":{"b"},   #a
    "b":{"c"},   #b
    "c":{"a","d","g"},   #c
    "d":{"e"},   #d
    "e":{"f"},   #e
    "f":{"d"},   #f
    "g":{"h"},   #g
    "h":{"i"},   #h
    "i":{"g"}    #i
}
'''

#tarjan算法求图中强连通分量
def tarjan(edges):
    index_counter = [0]
    stack = []
    lowlink = {}
    index = {}
    result = []
    
    def strongconnect(node):
        index[node] = index_counter[0]
        lowlink[node] = index_counter[0]
        index_counter[0] += 1
        stack.append(node)

        for v in edges[node]:
            if v not in lowlink:
                strongconnect(v)
                lowlink[node] = min(lowlink[node], lowlink[v])
            elif v in stack:
                lowlink[node] = min(lowlink[node], index[v])

        if lowlink[node] == index[node]:
            connected_component = []
            while True:
                v = stack.pop()
                connected_component.append(v)
                if v == node:
                    break
            result.append(connected_component)

    for node in edges:
        if node not in lowlink:
            strongconnect(node)

    return result

#计算强连通分量
import sys  # 导入sys模块
sys.setrecursionlimit(30000)  # 将默认的递归深度修改为3000
result = tarjan(idiom_jielong_graph)

#获取大的强连通分量
cercle = None
for key in result:
    if len(key) > 100:
        cercle = key
print(len(cercle))

#将cercle单独构建一个图，然后进行深度遍历
graph = {}

head_dict = {}
tail_dict = {}

for idiom in cercle:
    idiom_word_list = list(idiom)

    first_word = idiom_word_list[0]
    if first_word not in head_dict:
        head_dict[first_word] = [idiom]
    else:
        head_dict[first_word].append(idiom)

    last_word = idiom_word_list[-1]
    if last_word not in tail_dict:
        tail_dict[last_word] = [idiom]
    else:
        tail_dict[last_word].append(idiom)

for idiom in cercle:
    idiom_word_list = list(idiom)

    last_word = idiom_word_list[-1]

    if idiom not in graph:
        graph[idiom] = list()

        if last_word in head_dict:
            for key in head_dict[last_word]:
                graph[idiom].append(key)

#随机遍历，并对遍历结果进行修正，获得最长的成语接龙
def get_jielong_random(graph, output_file):
    #dfs 遍历 图
    def dfs(graph, node, visited):
        visited.append(node)
        for neighbor in graph[node]:
            if neighbor not in visited:
                dfs(graph, neighbor, visited)

    visited = list()
    #随机打散graph的key
    graph_keys = list(graph.keys())
    random.shuffle(graph_keys)
    for node in graph_keys:
        if node not in visited:
            dfs(graph, node, visited)

    #将初始遍历结果写入文件
    f=open("old_" + output_file,"w")
    f.writelines('\n'.join(visited))
    f.close()

    #修正visited
    new = []
    #指向 visited 的指针，初始为0
    index = 0

    max_len = 0
    max_index = None
    max_new = []
    while (index < len(visited) - 1):
        first = visited[index]
        second = visited[index + 1]

        #先将first加入new
        new.append(first)

        #检查first是否与second形成接龙
        if(list(first)[-1] != list(second)[0]):
            #new 向后看
            min_index = len(new) - 1
            for i in range(len(new) - 1, -1, -1):
                word = new[i]

                if list(word)[-1] == list(second)[0]:
                    min_index = i
                    break
            #获取向后看损失的长度
            backword_len = len(new) - 1 - min_index

            #visited 向前看
            max_index = index + 1
            for i in range(index + 1, len(visited), 1):
                word = visited[i]

                if list(first)[-1] == list(word)[0]:
                    max_index = i
                    break
            #获取向前看损失的长度
            forward_len = max_index - (index + 1)

            #如果前后看都失败，则跳出循环
            if backword_len == 0 and forward_len == 0:
                break

            #向后看失败，但是向前看成功了，向前跳过一部分
            elif backword_len == 0:
                new.append(visited[max_index])
                index = max_index
                continue

            #向前看失败，但是向后看成功了，向后丢弃一部分
            elif forward_len == 0:
                for i in range(backword_len):
                    new.pop()

            else:
                #向前向后看都成功了，但是向前看损失更少，向前跳过一部分
                if forward_len < backword_len:
                    new.append(visited[max_index])
                    index = max_index
                    continue

                #向前向后看都成功了，但是向后看损失更少，向后丢弃一部分
                else:
                    for i in range(backword_len):
                        new.pop()

        #end if (list(first)[-1] != list(second)[0]):

        index = index + 1

        if len(new) > max_len:
            max_len = len(new)
            max_new = new.copy()
            max_index = index - 1

    #end while (index < len(visited) - 1):

    #将修正后的结果写入文件
    f=open("new_" + output_file,"w")
    f.writelines('\n'.join(max_new))
    f.close()

    return max_index, max_len, max_new
#end def get_jielong_random(graph, output_file):


# #随机遍历，并对遍历结果进行修正，获得最长的成语接龙
# max_index, max_len, max_new = get_jielong_random(graph, output_file='output.txt')
# print(max_index, max_len, max_new[:10])

#将结果写入word文件
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from pypinyin import pinyin, lazy_pinyin, Style

#获取字体长度
from PIL import Image, ImageFont, ImageDraw
def get_text_length(text, font_size, font_path = "font/msyhbd.ttf"):
    font = ImageFont.truetype(font_path, font_size)
    image = Image.new('RGB', (1, 1))
    draw = ImageDraw.Draw(image)
    return draw.textlength(text, font=font)

'''
text = "yí zūn jiù jiào ， jiāo yī shí bǎi"
font_size = 12
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "移樽就教，教一识百"
font_size = 44
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "huí tiān yùn dǒu ， dǒu zhuǎn xīng yí"
font_size = 22
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")

text = "回天运斗，斗转星移"
font_size = 44
length = get_text_length(text, font_size)
print(f"The length of '{text}' in Arial {font_size} is {length} pixels.")
'''

def add_title(document, title):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(72)
    run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

#将一行成语加入word中，并对其进行拼音标注
def add_text_paragraph(document, text, text_font_size=44):
    #获得文字对应的拼音(方式一，通过lazy_pinyin获取)
    pinyin_text = ' '.join(lazy_pinyin(text, style=Style.TONE))

    #获得文字对应的拼音(方式二，通过pinyin_dict查找)
    idiom_list = text.split('，')
    pinyin_text = pinyin_dict[idiom_list[0]] + '，' + pinyin_dict[idiom_list[1]]

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    #设置段后间距（单倍行距，段后间距为0）
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_after = Pt(0)

    #固定文字字号为text_font_size(默认为44），推算拼音的字号大小，使得拼音与文字的长度保持一致
    text_length = get_text_length(text, text_font_size)

    #寻找最佳字号，使得拼音的长度与文字的长度保持一致
    pinyin_font_size = int(text_font_size / 2)
    min_gap = text_length
    min_pinyin_font_size = pinyin_font_size
    for size in range(pinyin_font_size - 10, pinyin_font_size + 10, 1):
        pinyin_length = get_text_length(pinyin_text, size)

        if abs(pinyin_length - text_length) < min_gap:
            min_gap = abs(pinyin_length - text_length)
            min_pinyin_font_size = size

        # print(size, abs(pinyin_length - text_length))

    #使用间距最小的字号作为拼音的字号
    pinyin_font_size = min_pinyin_font_size

    #写入拼音行
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(pinyin_text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(pinyin_font_size)

    #写入文本行
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(text_font_size)

def jielong_list_to_docx(jielong, outputfile, title="成语接龙"):
    document = Document()

    title = '成语接龙'
    add_title(document, title)

    for i in range(1, len(jielong), 2):
        text = jielong[i-1] + '，' + jielong[i]
        add_text_paragraph(document, text)

        #处理奇数情况的最后一个成语
        if i == len(jielong) - 2:
            text = jielong[i+1]
            add_text_paragraph(document, text)

    document.save(outputfile)

#多次随机遍历，并输出每次遍历的时间
for i in range(100):
    print (i, ' start \n')

    import datetime
    starttime = datetime.datetime.now()
    
    max_index, max_len, max_new = get_jielong_random(graph, output_file='output_' + str(i) + '.txt')

    endtime = datetime.datetime.now()

    print (i, ' end \n')
    print('cost seconds: ', (endtime - starttime).seconds, '\n')
    print(max_index, max_len, max_new[:10], '\n\n')

    #将最长的成语接龙写入word
    jielong_list_to_docx(max_new, 'max_new_' + str(i) + '_len_' + str(len(max_new)) + '.docx', title='成语接龙')

#验证new确实是连续的成语接龙
test = []
num = []
for i in range(1, len(new) - 2):
    first = new[i-1]
    second = new[i]

    test.append(first)

    #发现不一致，将不一致的从new中剔除即可
    if list(first)[-1] != list(second)[0]:
        num.append(i)
print(len(num))


#手工延长接龙尝试
last_word = list(new[-1])[-1]
while True:
    if last_word not in head_dict:
        break

    if len(head_dict[last_word]) == 0:
        break

    flag = 0
    for idiom in head_dict[last_word]:
        if idiom not in new:
            new.append(idiom)
            last_word = list(idiom)[-1]
            flag = 1
            break

    #没有找到合适的后续接龙成语，则跳出循环
    if flag == 0:
        break

first_word = list(new[0])[0]
while True:
    if first_word not in tail_dict:
        break

    if len(tail_dict[first_word]) == 0:
        break

    flag = 0
    for idiom in tail_dict[first_word]:
        if idiom not in new:
            new = [idiom] + new
            first_word = list(idiom)[0]
            flag = 1
            break

    #没有找到合适的后续接龙成语，则跳出循环
    if flag == 0:
        break


